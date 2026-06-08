import streamlit as st
from openai import OpenAI
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from datetime import datetime
import requests

# =========================
# 页面配置
# =========================
st.set_page_config(
    page_title="AI Marketing Copilot",
    page_icon="🚀",
    layout="wide"
)

# =========================
# Session状态
# =========================
if "history" not in st.session_state:
    st.session_state.history = []

# =========================
# Sidebar
# =========================
with st.sidebar:
    st.title("⚙️ 系统设置")
    model_name = st.selectbox(
        "模型选择",
        ["deepseek-chat", "deepseek-reasoner"]
    )
    template_type = st.selectbox(
        "内容模板",
        ["营销方案", "短视频脚本", "直播带货脚本", "客服回复"]
    )
    st.markdown("---")
    st.success(f"当前模型：{model_name}")
    st.markdown("### 历史记录")
    if len(st.session_state.history) == 0:
        st.caption("暂无历史记录")
    for item in reversed(st.session_state.history[-10:]):
        st.caption(item)

# =========================
# 页面输入
# =========================
st.title("🚀 AI Marketing Copilot")
st.caption("企业级AI营销内容生成平台")

col1, col2 = st.columns(2)
with col1:
    product_name = st.text_input("产品名称")
    target_customer = st.text_input("目标客户")
with col2:
    platform = st.selectbox("推广平台", ["小红书", "抖音", "Temu", "亚马逊"])

features = st.text_area(
    "产品特点",
    height=150,
    placeholder="316不锈钢\n保温24小时\n防漏设计"
)

# =========================
# Prompt模板
# =========================
prompt_templates = {
    "营销方案": """
输出：

# 商品标题
3个

# 核心卖点
3个

# 小红书文案

# 短视频口播稿

# 亚马逊卖点
5条
""",
    "短视频脚本": """
输出：

# 视频标题

# 分镜脚本

# 口播稿

# 拍摄建议
""",
    "直播带货脚本": """
输出：

# 开场话术

# 产品介绍

# 成交逼单话术

# 结束话术
""",
    "客服回复": """
输出：

# 售前回复

# 售后回复

# 差评回复

# 催付回复
"""
}

# =========================
# 生成内容
# =========================
if st.button("🚀 AI生成内容", use_container_width=True):
    if not product_name:
        st.warning("请输入产品名称")
        st.stop()

    try:
        api_key = st.secrets["DEEPSEEK_API_KEY"]
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

        prompt = f"""
你是一位资深电商运营专家。

产品名称：
{product_name}

目标客户：
{target_customer}

推广平台：
{platform}

产品特点：
{features}

任务类型：
{template_type}

要求：

{prompt_templates[template_type]}

请使用Markdown格式输出。
"""
        with st.spinner("AI正在思考中..."):
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.8
            )
            result = response.choices[0].message.content

        st.success("✅ 内容生成成功")
        st.markdown(result)

        # =========================
        # 飞书自动保存
        # =========================
        try:
            app_id = st.secrets["FEISHU_APP_ID"]
            app_secret = st.secrets["FEISHU_APP_SECRET"]
            base_id = st.secrets["BASE_ID"]
            table_id = st.secrets["TABLE_ID"]

            auth = requests.post(
                "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
                json={"app_id": app_id, "app_secret": app_secret}
            ).json()
            token = auth["tenant_access_token"]

            headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
            payload = {
                "fields": {
                    "产品名称": product_name,
                    "目标客户": target_customer,
                    "产品特点": features,
                    "AI文案": result
                }
            }

            url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{base_id}/tables/{table_id}/records"
            feishu_resp = requests.post(url, headers=headers, json=payload)

            if feishu_resp.status_code == 200:
                st.success("✅ 已自动保存到飞书")
            else:
                st.warning(f"飞书保存失败：{feishu_resp.text}")
        except Exception as e:
            st.warning(f"飞书保存失败：{e}")

        # =========================
        # Token统计
        # =========================
        try:
            input_tokens = response.usage.prompt_tokens
            output_tokens = response.usage.completion_tokens
            total_tokens = response.usage.total_tokens
        except:
            input_tokens = output_tokens = total_tokens = 0

        st.divider()
        c1, c2, c3 = st.columns(3)
        c1.metric("输入Token", input_tokens)
        c2.metric("输出Token", output_tokens)
        c3.metric("总Token", total_tokens)

        # =========================
        # 历史记录
        # =========================
        st.session_state.history.append(f"{datetime.now().strftime('%Y-%m-%d %H:%M')} | {product_name}")

        # =========================
        # Word导出
        # =========================
        doc = Document()
        doc.add_heading("AI营销方案", level=1)
        doc.add_paragraph(result)
        word_file = "marketing_plan.docx"
        doc.save(word_file)
        with open(word_file, "rb") as file:
            st.download_button("📄 下载Word", file, file_name=word_file)

        # =========================
        # Markdown导出
        # =========================
        st.download_button("📥 下载Markdown", result, file_name="marketing_plan.md")

        # =========================
        # PDF导出
        pdf_file = "marketing_plan.pdf"
        pdf = SimpleDocTemplate(pdf_file)
        styles = getSampleStyleSheet()
        content = [Paragraph(result.replace("\n", "<br/>"), styles["BodyText"])]
        pdf.build(content)
        with open(pdf_file, "rb") as file:
            st.download_button("📕 下载PDF", file, file_name=pdf_file)

    except Exception as e:
        st.error(f"发生错误：{e}")